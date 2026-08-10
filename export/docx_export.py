"""Word export of an AI Assistant answer, mirroring the PDF export's contents
(Task C4).

Like pdf_export, returns bytes rather than writing to disk so Streamlit's
st.download_button can serve them directly.
"""

from __future__ import annotations

import datetime
import io
import re
from typing import Any

import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt, RGBColor

from viz.charts import figure_to_image_bytes

MAX_TABLE_ROWS = 15


def _add_markdown_paragraph(document: Document, markdown_text: str) -> None:
    """Add narrative text, honouring the small Markdown subset the narrative
    prompts produce (**bold** and `-` bullets). Not a general Markdown parser --
    just enough that bold figures don't render as literal asterisks.
    """
    for raw_line in markdown_text.strip().split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        is_bullet = bool(re.match(r"^[-*]\s+", line))
        if is_bullet:
            line = re.sub(r"^[-*]\s+", "", line)
            paragraph = document.add_paragraph(style="List Bullet")
        else:
            paragraph = document.add_paragraph()

        # Split on bold spans, keeping the delimiters' contents.
        for segment in re.split(r"(\*\*.+?\*\*)", line):
            if not segment:
                continue
            run = paragraph.add_run(segment[2:-2] if segment.startswith("**") else segment)
            run.bold = segment.startswith("**")


def _add_meta_line(document: Document, text: str) -> None:
    """Add a small grey metadata line."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)


def build_answer_docx(
    question: str,
    narrative: str,
    sql: str | None = None,
    result: pd.DataFrame | None = None,
    figure: go.Figure | None = None,
    dataset_metadata: dict[str, Any] | None = None,
    applied_filters: dict[str, Any] | None = None,
) -> bytes:
    """Build a Word report for one AI Assistant answer and return its bytes.
    Same optional-section contract as build_answer_pdf().
    """
    document = Document()
    document.add_heading("AI Analytics Report", level=0)
    _add_meta_line(document, f"Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    if dataset_metadata:
        _add_meta_line(document, " | ".join(f"{k}: {v}" for k, v in dataset_metadata.items()))
    if applied_filters:
        _add_meta_line(document, "Filters - " + " | ".join(f"{k}: {v}" for k, v in applied_filters.items()))

    document.add_heading("Question", level=2)
    document.add_paragraph(question)

    document.add_heading("Answer", level=2)
    _add_markdown_paragraph(document, narrative)

    if figure is not None:
        document.add_heading("Chart", level=2)
        image_bytes = figure_to_image_bytes(figure, fmt="png", scale=2)
        document.add_picture(io.BytesIO(image_bytes), width=Inches(6.2))

    if result is not None and not result.empty:
        document.add_heading("Result data", level=2)
        shown = result.head(MAX_TABLE_ROWS)
        table = document.add_table(rows=1, cols=len(shown.columns))
        table.style = "Light Grid Accent 1"
        for i, column_name in enumerate(shown.columns):
            table.rows[0].cells[i].text = str(column_name)
        for row in shown.itertuples(index=False):
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        if len(result) > MAX_TABLE_ROWS:
            _add_meta_line(document, f"Showing first {MAX_TABLE_ROWS} of {len(result):,} rows.")

    if sql:
        document.add_heading("Generated SQL", level=2)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(sql)
        run.font.name = "Consolas"
        run.font.size = Pt(8)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
